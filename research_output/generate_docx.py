# -*- coding: utf-8 -*-
"""Build the APA-styled .docx from draft.md and render a landscape figure."""
import os, re
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
os.makedirs(ASSETS, exist_ok=True)

# ---------- Figure 1: research landscape timeline ----------
def make_figure():
    threads = [
        ("LLM planners", [("SayCan", 2022), ("Code as Policies", 2023),
                           ("ChatGPT for Robotics", 2024), ("LLVM-drone", 2025)]),
        ("VLA end-to-end", [("RT-1", 2022), ("RT-2", 2023), ("OpenVLA", 2024), ("VLA survey", 2025)]),
        ("Diffusion policy", [("Diffusion Policy", 2023)]),
        ("Sim-to-real RL", [("High-Speed Flight", 2021), ("Agile benchmark", 2022),
                            ("Swift", 2023), ("OmniDrones", 2024), ("Aerial Gym", 2025)]),
        ("Aerial VLN", [("AerialVLN", 2023), ("OpenFly", 2025),
                        ("CityNavAgent", 2025), ("UAV-VLN survey", 2026)]),
    ]
    colors = ["#2563eb", "#0d9488", "#9333ea", "#dc2626", "#ea580c"]
    fig, ax = plt.subplots(figsize=(9.2, 4.4))
    for i, (name, papers) in enumerate(threads):
        y = len(threads) - i
        years = [p[1] for p in papers]
        ax.plot([min(years), max(years)], [y, y], color=colors[i], lw=1.2, alpha=0.5, zorder=1)
        seen = {}
        for label, yr in papers:
            dy = 0.0
            seen[yr] = seen.get(yr, 0) + 1
            if seen[yr] > 1:
                dy = 0.16 * (seen[yr] - 1)
            ax.scatter(yr, y + dy, s=46, color=colors[i], zorder=3)
            ax.annotate(label, (yr, y + dy), textcoords="offset points",
                        xytext=(0, 7), ha="center", fontsize=7.2, color="#111827")
        ax.text(2020.55, y, name, ha="right", va="center", fontsize=9, fontweight="bold", color=colors[i])
    ax.set_xlim(2018.6, 2026.7)
    ax.set_ylim(0.3, len(threads) + 1.1)
    ax.set_yticks([])
    ax.set_xticks(range(2021, 2027))
    ax.set_xlabel("Year", fontsize=9)
    ax.set_title("Figure 1. Research landscape of embodied intelligence for UAVs (2021–2026)",
                 fontsize=10.5, fontweight="bold")
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    ax.tick_params(axis="x", labelsize=8)
    plt.tight_layout()
    out = os.path.join(ASSETS, "landscape.png")
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out

FIG = make_figure()

# ---------- Inline markdown (**bold**, *italic*) ----------
TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
def add_runs(paragraph, text):
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)

# ---------- Parse draft.md ----------
with open(os.path.join(HERE, "draft.md"), encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

i = 0
in_refs = False
n = len(lines)
while i < n:
    line = lines[i].rstrip()
    if not line.strip():
        i += 1; continue

    # Markdown table block
    if line.lstrip().startswith("|"):
        block = []
        while i < n and lines[i].lstrip().startswith("|"):
            block.append(lines[i].strip()); i += 1
        rows = [[c.strip() for c in r.strip("|").split("|")] for r in block
                if not re.match(r"^\|[\s:|-]+\|$", r)]
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    p = t.cell(ri, ci).paragraphs[0]
                    add_runs(p, cell)
                    for run in p.runs:
                        run.font.size = Pt(8)
                        if ri == 0:
                            run.bold = True
        continue

    if line.startswith("# "):
        h = doc.add_heading(line[2:], level=0)
        i += 1; continue
    if line.startswith("## "):
        title = line[3:]
        in_refs = title.strip().lower() == "references"
        doc.add_heading(title, level=1)
        i += 1; continue
    if line.startswith("### "):
        doc.add_heading(line[4:], level=2)
        i += 1; continue

    # Reference entry -> hanging indent
    if in_refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(6)
        add_runs(p, line)
        i += 1; continue

    # Insert Figure 1 right before the Synthesis section's first paragraph
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

# Embed figure near the end (after body, before nothing critical): add a figure page
doc.add_heading("Figure", level=1)
doc.add_picture(FIG, width=Inches(6.3))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap.add_run("Figure 1. Research landscape of embodied intelligence for UAVs (2021–2026).")
r.italic = True; r.font.size = Pt(9)

out_path = os.path.abspath(os.path.join(HERE, "Research_Report.docx"))
doc.save(out_path)
print("SAVED:", out_path)
print("FIGURE:", FIG)
